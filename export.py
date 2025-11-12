from __future__ import annotations
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def save_markdown_outputs(summary_output: str, assets: Dict[str, str], output_dir: str = "outputs") -> Path:
    """
    Save all generated marketing assets (summary + channel content) into a single Markdown file.

    Args:
        summary_output (str): The unified launch brief or summarised text.
        assets (Dict[str, str]): A dictionary containing channel outputs, e.g.
                                 {"linkedin": "...", "newsletter": "...", "blog": "..."}.
        output_dir (str): Directory to save the Markdown file (default: "outputs").

    Returns:
        Path: Path object of the saved Markdown file.
    """

    # Ensure output directory exists
    out_dir = Path(output_dir)
    out_dir.mkdir(exist_ok=True)

    # Build filename with timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    md_output = out_dir / f"GenieAI_Marketing_Outputs_{timestamp}.md"

    # Build Markdown content
    content = f"""# Genie AI – Feature Launch Marketing Outputs
_Generated: {datetime.now():%Y-%m-%d %H:%M}_

---

## Unified Launch Brief
{summary_output}

---

## LinkedIn Post
{assets.get('linkedin', '⚠️ LinkedIn post not found.')}

---

## Newsletter Email
{assets.get('newsletter', '⚠️ Newsletter email not found.')}

---

## Blog Post
{assets.get('blog', '⚠️ Blog post not found.')}
"""

    # Write file
    md_output.write_text(content, encoding="utf-8")

    print(f"✅ Saved Markdown to {md_output.resolve()}")
    return md_output


if __name__ == "__main__":
    # Example usage (for standalone testing)
    dummy_assets = {
        "linkedin": "LinkedIn example post here.",
        "newsletter": "Newsletter content example here.",
        "blog": "Blog post example here."
    }
    save_markdown_outputs("Example summary text.", dummy_assets)


# --- Markdown → DOCX helpers -------------------------------------------------

HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_PATTERN = re.compile(r"^[-*+]\s+(.*)$")
NUMBER_PATTERN = re.compile(r"^\d+[.)]\s+(.*)$")


def markdown_to_docx(markdown_text: str, output_path: Path | str, *, title: Optional[str] = None) -> Path:
    """Convert Markdown text into a simple Word document."""

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    in_code_block = False
    code_buffer: list[str] = []

    for raw_line in markdown_text.replace("\r\n", "\n").split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                para = doc.add_paragraph("\n".join(code_buffer))
                para.style = "Intense Quote"
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                code_buffer.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(raw_line)
            continue

        if not stripped:
            doc.add_paragraph()
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            doc.add_heading(heading_match.group(2), level=level)
            continue

        bullet_match = BULLET_PATTERN.match(stripped)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style="List Bullet")
            continue

        number_match = NUMBER_PATTERN.match(stripped)
        if number_match:
            doc.add_paragraph(number_match.group(1), style="List Number")
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph(stripped.lstrip("> "))
            para.style = "Intense Quote"
            continue

        doc.add_paragraph(line)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def markdown_file_to_docx(markdown_path: Path | str, output_path: Path | str | None = None) -> Path:
    """Read a Markdown file and save it as DOCX (same stem by default)."""

    src = Path(markdown_path)
    if output_path is None:
        output_path = src.with_suffix(".docx")
    text = src.read_text(encoding="utf-8")
    return markdown_to_docx(text, output_path, title=src.stem.replace("_", " "))
